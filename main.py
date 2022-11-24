from docxcompose.composer import Composer
from docx import Document
from docx.shared import Inches, Cm
import shutil
import os
from docx.shared import Pt
import codecs
import  random

def savePas(path, files):
    with codecs.open(path, 'w',encoding='utf-8-sig') as wfd:
        for f in files:
            with codecs.open(f, 'r',encoding='utf-8-sig') as fd:
                for line in fd.readlines():
                    line_encode = line.encode("utf-8", "ignore")
                    line_decode = line_encode.decode()
                    wfd.write(str(line_decode))
                wfd.write("\n")

print("Название папки с заданиями: ")
master_dir = os.curdir+'/'+input()+'/'
#master_dir = os.curdir+'/Контрольная 3/'
print("Номер работы: ")
test_num = input()
print("Кол-во вариантов: ")
#var_num = 12
var_num = int(input())

ttask_start = ''
ttask_end = ''
test_start = ''
test_end = ''
helper = ''

tasks_dict = {}

files = os.listdir(master_dir)
for file in files:
    if file == "Базовые файлы":
        base_dir = master_dir+"Базовые файлы"+"/"
        base_files = os.listdir(base_dir)
        for base_file in base_files:
            if base_file.startswith("TTask"+test_num):
                if base_file.endswith('start.pas'):
                    ttask_start = base_dir+base_file
                if base_file.endswith('end.pas'):
                    ttask_end = base_dir+base_file
            if base_file.startswith("TTask"+test_num+"Unittests"):
                if base_file.endswith('start.pas'):
                    test_start = base_dir + base_file
                if base_file.endswith('end.pas'):
                    test_end = base_dir + base_file
            if base_file.startswith("TestHelpers"):
                helper = base_dir+base_file
    if file.startswith("Задание"):
        task_num = file.split(' ')[1]
        tasks_dict[task_num] = []
        task_dir = master_dir +'/'+file+'/'
        tasks = os.listdir(task_dir)
        for task in tasks:
            words = task.split('_')
            if words[1] not in tasks_dict[task_num]:
                tasks_dict[task_num].append(words[1])

task_arrays = {}
for key in tasks_dict.keys():
    value = tasks_dict[key]
    koef = (var_num // len(value))
    if (var_num % len(value) != 0):
        koef += 1
    task_arrays[key] = value*koef
    random.shuffle(task_arrays[key])

output_dir = os.curdir+'/output/'
os.mkdir(output_dir)

for i in range(0,var_num):
    var = i+1
    var_dir = output_dir+"Вариант "+str(var)+"/"
    os.mkdir(var_dir)
    if helper != '':
        shutil.copy(helper, var_dir+"TestHelpers.pas")

    doc_filename = var_dir+'Вариант '+str(var)+'.docx'
    master = Document()
    style = master.styles['Normal']
    style.font.name = 'Arial'
    run = master.add_paragraph().add_run('Вариант '+str(var))
    sections = master.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    run.font.size = Pt(24)
    run.font.name = 'Arial'
    run.bold = True
    composer = Composer(master)

    unit_files = [ttask_start]
    test_files = [test_start]

    for key in task_arrays.keys():
        task_var = task_arrays[key][i]
        task_dir = master_dir + "Задание " + str(key) +"/"
        task_file =  task_dir + str(key)+"_"+str(task_var)+"_task.docx"
        test_file = task_dir + str(key) + "_" + str(task_var) + "_test.pas"
        unit_file = task_dir + str(key) + "_" + str(task_var) + "_unit.pas"

        if os.path.exists(unit_file):
            unit_files.append(unit_file)
        if os.path.exists(test_file):
            test_files.append(test_file)
        doc = Document(task_file)
        composer.append(doc)

    unit_files.append(ttask_end)
    test_files.append(test_end)
    savePas(var_dir + "TTask" + test_num + ".pas",unit_files)
    savePas(var_dir + "TTask3Unittests" + test_num + ".pas", test_files)

    composer.save(doc_filename)


"""
master = Document("master.docx")
composer = Composer(master)
doc1 = Document("doc1.docx")
composer.append(doc1)
composer.save("combined.docx")
"""

"""
with open('output_file.txt','wb') as wfd:
    for f in ['seg1.txt','seg2.txt','seg3.txt']:
        with open(f,'rb') as fd:
            shutil.copyfileobj(fd, wfd)
"""